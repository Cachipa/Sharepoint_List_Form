from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory
from shareplum import Site
from shareplum import Office365
from docx import Document
import os

app = Flask(__name__)
app.secret_key = "your_secret_key"  # Necessário para usar flash messages

# SharePoint credentials
username = None
password = None


@app.route("/", methods=["GET", "POST"])
def login():
    global username, password
    if request.method == "POST":
        username = request.form.get("email")
        password = request.form.get("password")

        if not username or not password:
            flash("Email e senha não podem estar vazios!", "error")
            return redirect(url_for("login"))

        try:
            # Attempt to authenticate with SharePoint
            Office365('https://meioambientemg.sharepoint.com', username=username, password=password).GetCookies()
            flash("Login bem-sucedido!", "success")
            return redirect(url_for("main"))
        except Exception as e:
            flash(f"Erro no login: {e}", "error")
            return redirect(url_for("login"))

    return render_template("login.html")


@app.route("/main", methods=["GET", "POST"])
def main():
    try:
        # Authenticate and connect to SharePoint
        authcookie = Office365('https://meioambientemg.sharepoint.com', username=username, password=password).GetCookies()
        site = Site('https://meioambientemg.sharepoint.com/sites/BasedeDados', authcookie=authcookie)
        sp_list = site.List('Base de Dados')

        # Inicializar filtros
        status_filter = request.form.get("status_filter") if request.method == "POST" else None
        id_filter = request.form.get("id_filter") if request.method == "POST" else None

        # Construir a query com base nos filtros
        query = {}
        if status_filter:
            query.setdefault('Where', []).append(('Eq', 'Status', status_filter))
        if id_filter:
            query.setdefault('Where', []).append(('Eq', 'ID', id_filter))

        # Buscar itens com base na query
        items = sp_list.GetListItems(query=query) if query else sp_list.GetListItems()
    except Exception as e:
        flash(f"Erro ao acessar a lista do SharePoint: {e}", "error")
        items = []

    return render_template("main.html", items=items, status_filter=status_filter, id_filter=id_filter)


@app.route("/form", methods=["GET", "POST"])
def form():
    if request.method == "POST":
        # Get form data
        form_data = {
            'ID': item_id,
            'Status': request.form.get("status"),
            'Numero SEI': request.form.get("numero_sei"),
            'Nome': request.form.get("nome"),
            'Endereço': request.form.get("endereco"),
            'CPF/CNPJ': request.form.get("cpf_cnpj"),
            'Endereço Numero': request.form.get("endereco_numero"),
            'Bairro': request.form.get("bairro"),
            'UF': request.form.get("uf"),
            'CEP': request.form.get("cep"),
            'Telefone': request.form.get("telefone"),
        }

        # Check for required fields
        if not form_data['Status'] or not form_data['Numero SEI']:
            flash("Os campos 'Status' e 'Número SEI' são obrigatórios!", "error")
            return redirect(url_for("form"))

        try:
            # Authenticate and connect to SharePoint
            authcookie = Office365('https://meioambientemg.sharepoint.com', username=username, password=password).GetCookies()
            site = Site('https://meioambientemg.sharepoint.com/sites/BasedeDados', authcookie=authcookie)
            sp_list = site.List('Base de Dados')

            # Insert the new item
            sp_list.UpdateListItems(data=[form_data], kind='New')
            flash("Item inserido com sucesso no SharePoint!", "success")
        except Exception as e:
            flash(f"Erro ao inserir item: {e}", "error")

    return render_template("form.html")


@app.route("/edit/<item_id>", methods=["GET", "POST"])
def edit(item_id):
    try:
        # Authenticate and connect to SharePoint
        authcookie = Office365('https://meioambientemg.sharepoint.com', username=username, password=password).GetCookies()
        site = Site('https://meioambientemg.sharepoint.com/sites/BasedeDados', authcookie=authcookie)
        sp_list = site.List('Base de Dados')

        print(f"Editando item com ID: {item_id}")

        if request.method == "POST":
            # Get updated form data
            form_data = {
                'ID': item_id,  # Ensure the ID is included for updating the item
                'Status': request.form.get("status"),
                'Numero SEI': request.form.get("numero_sei"),
                'Nome': request.form.get("nome"),
                'Endereço': request.form.get("endereco"),
                'CPF/CNPJ': request.form.get("cpf_cnpj"),
                'Endereço Numero': request.form.get("endereco_numero"),
                'Bairro': request.form.get("bairro"),
                'UF': request.form.get("uf"),
                'CEP': request.form.get("cep"),
                'Telefone': request.form.get("telefone"),
            }

            # Update the existing item
            sp_list.UpdateListItems(data=[form_data], kind='Update')
            flash("Item atualizado com sucesso no SharePoint!", "success")
            return redirect(url_for("main"))

        # Fetch the item data for pre-filling the form
        item = sp_list.GetListItems(
            fields=['ID', 'Status', 'Numero SEI', 'Nome', 'Endereço', 'CPF/CNPJ', 'Endereço Numero', 'Bairro', 'UF', 'CEP', 'Telefone'],
            query={'Where': [('Eq', 'ID', item_id)]}  # Use o filtro correto para o campo ID
        )
        if not item:
            flash("Item não encontrado!", "error")
            return redirect(url_for("main"))

        # Pass the item data to the form
        return render_template("form.html", item=item[0])
    except Exception as e:
        flash(f"Erro ao acessar ou atualizar o item: {e}", "error")
        return redirect(url_for("main"))


@app.route("/download/<item_id>", methods=["GET"])
def download(item_id):
    try:
        # Authenticate and connect to SharePoint
        authcookie = Office365('https://meioambientemg.sharepoint.com', username=username, password=password).GetCookies()
        site = Site('https://meioambientemg.sharepoint.com/sites/BasedeDados', authcookie=authcookie)
        sp_list = site.List('Base de Dados')

        print(f"Gerando download para o item com ID: {item_id}")

        # Fetch the item data
        item = sp_list.GetListItems(
            fields=['Nome', 'Endereço', 'Telefone'],
            query={'Where': [('Eq', 'ID', item_id)]}  # Use o filtro correto para o campo ID
        )
        if not item:
            flash("Item não encontrado!", "error")
            return redirect(url_for("main"))

        item = item[0]  # Get the first (and only) item

        # Load the Word template
        template_path = os.path.join("Site-1","2Modelo Parecer (Fabio) 2.docx")
        if not os.path.exists(template_path):
            flash("Template Word não encontrado!", "error")
            return redirect(url_for("main"))
        print(f"Caminho do template Word: {template_path}")
        doc = Document(template_path)

        # Função para substituir texto em parágrafos e células de tabelas
        def replace_text_in_paragraphs(paragraphs, replacements):
            for paragraph in paragraphs:
                for key, value in replacements.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, value)

        # Função para substituir texto em tabelas
        def replace_text_in_tables(tables, replacements):
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_text_in_paragraphs(cell.paragraphs, replacements)

        # Substituições a serem feitas
        replacements = {
            "{{Nome}}": item.get("Nome", ""),
            "{{Endereço}}": item.get("Endereço", ""),
            "{{Telefone}}": item.get("Telefone", "")
        }

        # Substituir texto nos parágrafos do corpo do documento
        replace_text_in_paragraphs(doc.paragraphs, replacements)

        # Substituir texto nas tabelas
        replace_text_in_tables(doc.tables, replacements)

        # Substituir texto nos cabeçalhos e rodapés
        for section in doc.sections:
            replace_text_in_paragraphs(section.header.paragraphs, replacements)
            replace_text_in_paragraphs(section.footer.paragraphs, replacements)

        # Save the filled document to a temporary file
        output_directory = os.path.abspath("Site-1/downloads")
        output_filename = f"output_{item_id}.docx"
        output_path = os.path.join(output_directory, output_filename)
        doc.save(output_path)

        if not os.path.exists(output_path):
            flash("Erro ao salvar o arquivo Word!", "error")
            return redirect(url_for("main"))

        print(f"Arquivo Word gerado em: {output_path}")

       # Enviar o arquivo para o navegador
        return send_from_directory(
            directory=output_directory,
            filename=output_filename,  # Use 'path' em vez de 'filename' para maior clareza
            as_attachment=True,    # Força o download
            attachment_filename=f"Parecer_{item_id}.docx"  # Nome do arquivo no download
        )

    except Exception as e:
        print(f"Erro ao gerar o arquivo Word: {e}", "error")
        return redirect(url_for("main"))


if __name__ == "__main__":
    app.run()